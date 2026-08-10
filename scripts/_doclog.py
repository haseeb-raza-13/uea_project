"""Shared helpers for the automatic Word-document logging mechanism.

Used by init_log.py (creates the log) and run_logged.py (appends to it).
Not meant to be run directly.
"""
from pathlib import Path

import docx
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = PROJECT_ROOT / "docs" / "QC_Project_Log.docx"

HEADERS = ["Step", "Timestamp", "Purpose", "Tool (version)", "Command", "Status", "Notes / Key Output"]


def load_doc():
    if not DOC_PATH.exists():
        raise FileNotFoundError(f"{DOC_PATH} does not exist yet — run init_log.py first.")
    return docx.Document(str(DOC_PATH))


def get_log_table(doc):
    return doc.tables[0]


def next_step_number(table):
    rows = table.rows
    if len(rows) <= 1:
        return 1
    try:
        return int(rows[-1].cells[0].text.strip()) + 1
    except ValueError:
        return len(rows)


def append_row(doc, timestamp, purpose, tool_version, command, status, notes):
    table = get_log_table(doc)
    step = next_step_number(table)
    row = table.add_row()
    values = [
        str(step),
        timestamp.strftime("%Y-%m-%d %H:%M:%S"),
        purpose,
        tool_version,
        command,
        status,
        notes,
    ]
    for cell, value in zip(row.cells, values):
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
    doc.save(str(DOC_PATH))
    return step
