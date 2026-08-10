#!/usr/bin/env python3
"""One-time initializer for docs/QC_Project_Log.docx.

Creates the log document and backfills the environment-setup steps that were
performed before this logging mechanism existed. Run once; safe to re-run
(will not overwrite an existing log).
"""
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

import _doclog as dl


def build_doc():
    doc = docx.Document()

    title = doc.add_heading("Sequencing QC & Analysis Log", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Acinetobacter baumannii short-read / long-read WGS project").italic = True

    meta = doc.add_paragraph()
    meta.add_run(f"Project path: {dl.PROJECT_ROOT}\n")
    meta.add_run(f"Log created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    meta.add_run(
        "Every setup and analysis step performed for this project is recorded below, "
        "in order, for later use in the Methods section of a manuscript."
    )

    doc.add_heading("Step-by-step log", level=1)
    table = doc.add_table(rows=1, cols=len(dl.HEADERS))
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, dl.HEADERS):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.save(str(dl.DOC_PATH))
    return doc


def main():
    if dl.DOC_PATH.exists():
        print(f"{dl.DOC_PATH} already exists — not overwriting.")
        return

    dl.DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()

    backfilled = [
        dict(
            purpose="Install a conda/mamba package manager inside WSL to host the sequencing QC tool environment",
            tool_version="Miniforge3 (conda 26.3.2, mamba 2.5.0)",
            command=(
                'curl -sL -o Miniforge3.sh '
                'https://github.com/conda-forge/miniforge/releases/latest/download/Miniforge3-Linux-x86_64.sh '
                '&& bash Miniforge3.sh -b -p "$HOME/miniforge3"'
            ),
            status="Success",
            notes="Installed to /home/haseeb/miniforge3 inside WSL Ubuntu-22.04.",
        ),
        dict(
            purpose="Configure package channels for bioinformatics tool installation",
            tool_version="conda 26.3.2",
            command=(
                "conda config --set channel_priority strict "
                "&& conda config --add channels conda-forge && conda config --add channels bioconda"
            ),
            status="Success",
            notes=(
                "Channel order: bioconda (top priority), conda-forge. "
                "Strict channel priority enabled (standard bioconda requirement)."
            ),
        ),
        dict(
            purpose="Create the dedicated QC/trimming tool environment for this project",
            tool_version="mamba 2.5.0",
            command="mamba create -y -n seqqc python=3.10 fastqc fastp trimmomatic cutadapt bbmap multiqc python-docx",
            status="Success",
            notes=(
                "Environment 'seqqc' created with: FastQC, fastp, Trimmomatic, cutadapt, "
                "BBMap (bbduk.sh), MultiQC, python-docx. Exact tool versions confirmed in the "
                "verification steps below."
            ),
        ),
    ]

    now = datetime.now()
    for entry in backfilled:
        dl.append_row(doc, timestamp=now, **entry)

    print(f"Created {dl.DOC_PATH} with {len(backfilled)} backfilled setup step(s).")


if __name__ == "__main__":
    main()
