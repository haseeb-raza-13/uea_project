#!/usr/bin/env python3
"""Generate docs/Assembly_Comparison_Report.docx: a unified, hierarchical,
step-by-step account of the hybrid-assembly-vs-reference comparison work
(FastANI + NucDiff) done for all 5 strains, including every figure and both
summary tables. Times New Roman throughout; headings 14pt, body 12pt.
"""
import csv
from pathlib import Path
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = PROJECT_ROOT / "docs" / "Assembly_Comparison_Report.docx"
COMP_ROOT = PROJECT_ROOT / "analysis" / "assembly_comparison"

FONT_NAME = "Times New Roman"
HEADING_SIZE = Pt(14)
BODY_SIZE = Pt(12)
CAPTION_SIZE = Pt(11)
BLACK = RGBColor(0, 0, 0)

STRAINS = [
    ("mff", "ATCC 17978-mff"),
    ("VU", "ATCC 17978-VU"),
    ("AB42", "AB042"),
    ("AB30", "AB030"),
    ("Lac4", "LAC-4"),
]

FRAGMENTED = {"AB30", "Lac4"}


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.font.color.rgb = BLACK


def add_heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE
    run.font.color.rgb = BLACK
    run.bold = True
    return h


def add_para(doc, runs, bullet=False, numbered=False):
    style = "List Bullet" if bullet else ("List Number" if numbered else None)
    p = doc.add_paragraph(style=style)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opts in runs:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = BODY_SIZE
        r.font.color.rgb = BLACK
        if opts.get("bold"):
            r.bold = True
        if opts.get("highlight"):
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if opts.get("italic"):
            r.italic = True
    return p


def B(text):
    return (text, {"bold": True})


def H(text):
    return (text, {"bold": True, "highlight": True})


def T(text):
    return (text, {})


def style_cell(cell, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = Pt(9.5)
            r.font.color.rgb = BLACK
            r.bold = bold


def load_summary():
    with open(COMP_ROOT / "summary_table.csv", newline="") as f:
        rows = {row["strain"]: row for row in csv.DictReader(f)}
    return rows


def add_ani_completeness_table(doc, summary):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Strain", "FastANI (%)", "Contigs", "Largest contig (%)", "Assembly length (bp)", "Complete?"]
    for cell, h in zip(table.rows[0].cells, headers):
        cell.text = h
        style_cell(cell, bold=True)
    for strain, display in STRAINS:
        r = summary[strain]
        row = table.add_row().cells
        row[0].text = display
        row[1].text = f"{float(r['ani_percent']):.4f}"
        row[2].text = r["n_contigs"]
        row[3].text = f"{float(r['largest_contig_frac']) * 100:.1f}%"
        row[4].text = f"{int(r['total_assembly_len']):,}"
        row[5].text = "No -- fragmented" if strain in FRAGMENTED else "Yes"
        for cell in row:
            style_cell(cell)
    doc.add_paragraph()


def add_variant_counts_table(doc, summary):
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    headers = ["Strain", "Total", "Substitutions", "Insertions", "Deletions", "Relocations", "Inversions", "Reshufflings (excluded)"]
    for cell, h in zip(table.rows[0].cells, headers):
        cell.text = h
        style_cell(cell, bold=True)
    for strain, display in STRAINS:
        r = summary[strain]
        row = table.add_row().cells
        row[0].text = display
        row[1].text = r["total_differences"]
        row[2].text = r["substitutions"]
        row[3].text = r["insertions"]
        row[4].text = r["deletions"]
        row[5].text = r["relocations"]
        row[6].text = r["inversions"]
        row[7].text = r["reshufflings"]
        for cell in row:
            style_cell(cell)
    doc.add_paragraph()


def add_figure(doc, image_path, fig_num, caption_text, width=6.5):
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = cap.add_run(f"Figure {fig_num}. ")
    label_run.font.name = FONT_NAME
    label_run.font.size = CAPTION_SIZE
    label_run.font.color.rgb = BLACK
    label_run.bold = True
    label_run.italic = True
    text_run = cap.add_run(caption_text)
    text_run.font.name = FONT_NAME
    text_run.font.size = CAPTION_SIZE
    text_run.font.color.rgb = BLACK
    text_run.italic = True


STRAIN_INTERPRETATION = {
    "mff": [T(
        "Near-identical to the public reference (FastANI 99.9996%); the hybrid assembly is "
        "complete (4 contigs, largest contig 95.7% of total length). All 58 detected differences "
        "are small-scale (43 substitutions, 9 insertions, 1 deletion, 2 relocations) and can be "
        "treated as high-confidence, genuine calls."
    )],
    "VU": [T(
        "Also near-identical to the reference (FastANI 99.9989%) with a complete assembly (2 "
        "contigs, largest 99.4%). 114 differences detected (102 substitutions, 4 insertions, 5 "
        "deletions, 1 relocation), all high-confidence given the assembly quality."
    )],
    "AB42": [T(
        "The cleanest comparison of the five (FastANI 99.9986%, 3 contigs, largest 99.4% of "
        "total length). 110 differences (101 substitutions, 3 insertions, 3 deletions, 1 "
        "relocation), all high-confidence."
    )],
    "AB30": [
        [H("This assembly is fragmented: 175 contigs, and the largest single contig covers only "
           "29.0% of total assembly length (vs. 95-99% for mff/VU/AB42)."),
         T(" FastANI is still high (99.9849%) because ANI is computed only over the fraction of "
           "the assembly that does align, but the 10,005 NucDiff differences reported for this "
           "strain (9,281 substitutions, 434 insertions, 232 deletions, 32 relocations, 2 "
           "inversions) should not be read at face value.")],
        [T(
            "A fragmented assembly means many of NucDiff's calls are artifacts of contig "
            "boundaries and lower-confidence, harder-to-assemble sequence rather than genuine "
            "biological differences from the reference. The 2 inversion calls are the most "
            "notable individual finding here -- see the case study in Section 7 -- but even "
            "those carry the same caveat and are not yet confirmed."
        )],
    ],
    "Lac4": [
        [H("The most fragmented assembly in this dataset: 192 contigs, largest only 6.1% of "
           "total assembly length -- no single contig captures even a tenth of the genome."),
         T(" FastANI is still 99.9309%, but the 11,701 reported differences (10,936 "
           "substitutions, 459 insertions, 248 deletions, 17 relocations, 0 inversions) should "
           "be treated as substantially inflated by assembly fragmentation, not as a genuine "
           "count of biological variation.")],
        [B("Recommendation: "), T(
            "re-run or QC the Unicycler hybrid assembly for AB30 and Lac4 (check read coverage, "
            "long-read quality/quantity, and the assembly graph) before drawing biological "
            "conclusions from their variant counts."
        )],
    ],
}


def build():
    doc = docx.Document()
    set_default_font(doc)
    summary = load_summary()

    title = doc.add_heading("", level=0)
    run = title.add_run("Hybrid Assembly vs. Reference Genome Comparison: Five Acinetobacter baumannii Strains")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, [
        T("Project: "), B("uea_project"), T(". Strains: "),
        B("mff, VU, AB42, AB30, Lac-4"), T(". Date: "),
        B(datetime.now().strftime("%Y-%m-%d")),
    ])
    add_para(doc, [T("Tools: "), B("FastANI 1.34"), T(" (Average Nucleotide Identity), "), B("NucDiff 2.0.3"),
                    T(" (SNPs and structural rearrangements, built on MUMmer/nucmer)."), ])
    add_para(doc, [T("Data referenced here: "), B("analysis/assembly_comparison/"), T(".")])

    # ---- 1. What was done today, step by step -------------------------------
    add_heading(doc, "1. What Was Done Today")
    add_para(doc, [T(
        "This report covers one continuous session comparing each strain's Unicycler hybrid "
        "assembly (long-read + short-read) against its publicly available reference genome. "
        "The steps below were carried out in order:"
    )])
    steps = [
        "Adapted the assembly-comparison workflow to run locally (no HPC/SLURM access was "
        "available), building tools/compare_assemblies.py + .sh to run FastANI and NucDiff for "
        "all 5 strains.",
        "Ran the comparison for all 5 strains, producing per-strain ANI values, SNP calls, and "
        "structural-variant calls under analysis/assembly_comparison/.",
        "Built tools/summarize_assembly_comparison.py to consolidate FastANI + NucDiff output "
        "into one cross-strain table (summary_table.csv), including assembly-completeness "
        "metrics (contig count, largest-contig fraction).",
        "Discussed a machine-learning-based variant confidence score (Variant Quality Score "
        "Recalibration, VQSR) as requested; determined it does not fit this data (no cohort, no "
        "truth set, wrong data type -- see Section 8) and, per direction, did not build it.",
        "Built whole-genome synteny dot plots per strain (Section 6, Figures 1/3/5/7/9). This "
        "surfaced an important finding: the AB30 and Lac4 assemblies are badly fragmented "
        "(Section 5), which was flagged before proceeding further.",
        "Built circular genome-wide variant ideograms per strain (Figures 2/4/6/8/10), "
        "correcting a labeling issue along the way: NucDiff's \"reshuffling\" calls reflect the "
        "assembly's arbitrary circular start point, not real rearrangement, and are excluded "
        "from the rearrangement counts shown.",
        "Built a cross-strain summary composite figure (Figure 11: variant-type counts, ANI, "
        "assembly completeness).",
        "Built a zoomed case-study figure for AB30's 2 candidate inversions (Figure 12), with "
        "caveats based on the assembly-quality evidence at that locus.",
        "Compiled all of the above into this report.",
    ]
    for s in steps:
        add_para(doc, [T(s)], numbered=True)

    # ---- 2. Methods -----------------------------------------------------------
    add_heading(doc, "2. Methods")
    add_para(doc, [T(
        "For each strain, the hybrid assembly (query) was compared against its public reference "
        "genome (reference) using two tools: "
    ), B("FastANI"), T(
        " for whole-genome Average Nucleotide Identity, and "
    ), B("NucDiff"), T(
        " (built on MUMmer/nucmer whole-genome alignment) for base-level SNPs/indels and "
        "larger structural differences (relocations, inversions, duplications, translocations). "
        "Both tools were installed in a dedicated conda environment (ani_nucdiff) and run "
        "locally; NucDiff's own \"reshuffling\" category, which for a circular chromosome "
        "reflects the assembly's and reference's different arbitrary start points rather than "
        "real rearrangement, is reported separately and excluded from rearrangement counts in "
        "the figures and tables below."
    )])

    # ---- 3. Cross-strain summary tables ---------------------------------------
    add_heading(doc, "3. Cross-Strain Summary Tables")
    add_para(doc, [T("From "), B("analysis/assembly_comparison/summary_table.csv"), T(
        " (itself built from summary_ani.tsv plus each strain's NucDiff output).")])
    add_heading(doc, "3.1 ANI and assembly completeness", level=2)
    add_ani_completeness_table(doc, summary)
    add_heading(doc, "3.2 NucDiff variant-type counts", level=2)
    add_variant_counts_table(doc, summary)

    # ---- 4. Key finding: fragmentation -----------------------------------------
    add_heading(doc, "4. Key Finding: Two of the Five Assemblies Are Fragmented")
    add_para(doc, [H(
        "AB30 (175 contigs, largest contig = 29.0% of total length) and Lac4 (192 contigs, "
        "largest contig = 6.1%) are substantially more fragmented than mff, VU, and AB42 (2-4 "
        "contigs each, largest contig 95-99%)."
    )])
    add_para(doc, [T(
        "This was discovered while building the dot plots (Section 6) and is the single most "
        "important caveat in this report: it means AB30's and Lac4's much higher NucDiff "
        "variant counts (Section 3.2) are substantially assembly-fragmentation artifacts, not "
        "confirmed biological differences from the reference. Every figure and table below "
        "involving AB30 or Lac4 carries this caveat explicitly."
    )])

    # ---- 5. Per-strain results --------------------------------------------------
    add_heading(doc, "5. Per-Strain Results")
    fig_num = 1
    for strain, display in STRAINS:
        add_heading(doc, f"5.{[s for s, _ in STRAINS].index(strain) + 1} Strain {display}", level=2)
        r = summary[strain]
        add_para(doc, [
            B("FastANI: "), T(f"{float(r['ani_percent']):.4f}%   "),
            B("Contigs: "), T(f"{r['n_contigs']}   "),
            B("Largest contig: "), T(f"{float(r['largest_contig_frac']) * 100:.1f}% of assembly   "),
            B("Total NucDiff differences: "), T(r["total_differences"]),
        ])
        interp = STRAIN_INTERPRETATION[strain]
        if interp and isinstance(interp[0], tuple):
            add_para(doc, interp)
        else:
            for line in interp:
                add_para(doc, line)

        dotplot_path = COMP_ROOT / strain / "figures" / f"{strain}_dotplot.png"
        circos_path = COMP_ROOT / strain / "figures" / f"{strain}_circos.png"
        add_figure(doc, dotplot_path, fig_num,
                   f"Strain {display}. Whole-genome synteny dot plot: hybrid assembly vs. reference, "
                   "colored by alignment orientation. oriC/ter marked.")
        fig_num += 1
        doc.add_paragraph()
        add_figure(doc, circos_path, fig_num,
                   f"Strain {display}. Circular genome-wide variant landscape: substitution density, "
                   "indel density, and structural rearrangements (fragmentation artifacts excluded).")
        fig_num += 1
        doc.add_page_break()

    # ---- 6. Cross-strain summary figure -----------------------------------------
    add_heading(doc, "6. Cross-Strain Summary Figure")
    add_para(doc, [T(
        "Figure 11 brings together the variant-type counts, ANI, and assembly-completeness "
        "metrics from Section 3 into a single composite, making the fragmentation-vs-variant-"
        "count relationship visually explicit: the two strains below the 90% completeness "
        "threshold (Panel C) are exactly the two strains with order-of-magnitude higher variant "
        "counts (Panel A)."
    )])
    add_figure(doc, COMP_ROOT / "figures_summary" / "assembly_comparison_summary.png", fig_num,
               "Cross-strain summary: NucDiff variant-type counts (log scale), FastANI, and "
               "assembly completeness, all 5 strains.")
    fig_num += 1
    doc.add_page_break()

    # ---- 7. Case study: AB30 inversions -----------------------------------------
    add_heading(doc, "7. Case Study: AB30's Candidate Inversions")
    add_para(doc, [T(
        "AB30 is the only strain with NucDiff-called inversions (2 total). Both sit on the same "
        "106,967bp assembly contig, not the 1.5Mb main chromosome scaffold, and local alignment "
        "identity at these breakpoints (97.4-97.9%) is below this comparison's genome-wide mean "
        "(99.76%). The two inversion blocks map to opposite ends of that same contig onto two "
        "reference windows only ~460bp apart -- a pattern more consistent with a contig-end/"
        "assembly-junction artifact than a clean internal inversion."
    )])
    add_para(doc, [H(
        "These should be treated as a candidate finding, not a confirmed rearrangement, pending "
        "validation against raw long-read data (checking for reads that span the breakpoint "
        "cleanly)."
    )])
    add_figure(doc, COMP_ROOT / "figures_locus" / "AB30_inversion_loci.png", fig_num,
               "Strain AB030. Candidate inversion breakpoints, reference (forward, top) vs. "
               "assembly contig (reverse, bottom); crossed ribbon is the standard synteny-plot "
               "signature of an inversion.")
    fig_num += 1
    doc.add_page_break()

    # ---- 8. Confidence scoring note ----------------------------------------------
    add_heading(doc, "8. Note: Variant Confidence Scoring (VQSR)")
    add_para(doc, [T(
        "A machine-learning-based confidence score for individual variant/rearrangement calls "
        "(GATK-style Variant Quality Score Recalibration) was discussed. VQSR fits a Gaussian-"
        "mixture model over per-site annotations from a population-scale probabilistic variant "
        "caller, calibrated against a known truth/training set (e.g. dbSNP, HapMap). This "
        "dataset -- 5 single-genome-vs-reference whole-genome alignments, with no cohort and no "
        "truth set, often only 60-100 variant records per (well-assembled) strain -- does not "
        "provide enough data to fit such a model, and no equivalent truth set exists for these "
        "bacterial strains. A scoped, appropriate alternative (alignment-quality filtering "
        "using signals already in the NucDiff/MUMmer output, or long-read breakpoint validation "
        "for structural calls specifically) was proposed; by direction, this was deferred and "
        "not built in this session."
    )])

    # ---- 9. Caveats and next steps -------------------------------------------------
    add_heading(doc, "9. Caveats and Recommended Next Steps")
    add_para(doc, [B("AB30 and Lac4 need re-assembly or assembly QC "),
                    T("before their variant counts or the AB30 inversions can be treated as biological findings, "
                      "not assembly artifacts (Section 4, Section 7).")], bullet=True)
    add_para(doc, [B("Variant confidence scoring "),
                    T("was scoped but not built (Section 8); revisit if a specific finding needs it.")], bullet=True)
    add_para(doc, [B("docs/QC_Project_Log.docx "),
                    T("was found corrupted (zero-byte content) partway through this session and could not be "
                      "repaired; today's steps were not logged there. Flagging for awareness -- let me know if "
                      "you'd like it reinitialized.")], bullet=True)
    add_para(doc, [B("Reshuffling calls "),
                    T("(whole-genome start-point offset between circular sequences) are excluded from "
                      "rearrangement counts throughout; this is a display/interpretation choice, not a data-quality "
                      "issue.")], bullet=True)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
