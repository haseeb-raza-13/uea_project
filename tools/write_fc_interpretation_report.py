#!/usr/bin/env python3
"""Generate docs/FoldChange_Interpretation.docx: an interpretation of the
Meropenem-vs-untreated log2 fold-change analysis (pooled overall test, its
transposed/horizontal variant, and the classic per-window volcano plot) for
all five strains. Times New Roman throughout; headings 14pt, body 12pt.

docs/ is not tracked by git (see .gitignore), so this report is intentionally
not committed -- only the analysis/ and tools/ pipeline that produces the
underlying figures and CSVs is.
"""
import csv
import sys
from pathlib import Path
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = PROJECT_ROOT / "docs" / "FoldChange_Interpretation.docx"
ANALYSIS_ROOT = PROJECT_ROOT / "analysis" / "replication_profile"

FONT_NAME = "Times New Roman"
HEADING_SIZE = Pt(14)
BODY_SIZE = Pt(12)
CAPTION_SIZE = Pt(11)
BLACK = RGBColor(0, 0, 0)

STRAINS = [
    ("mff", "ATCC 17978-mff"),
    ("VU", "ATCC 17978-VU"),
    ("AB30", "AB030"),
    ("AB42", "AB042"),
    ("Lac-4", "LAC-4"),
]

PHASE_LABEL = {"stat": "Stationary", "exp": "Exponential"}

# The phase highlighted (embedded figure + interpretation) for each strain,
# chosen as whichever phase carries the more informative signal for that strain.
HIGHLIGHT_PHASE = {"mff": "exp", "VU": "exp", "AB30": "stat", "AB42": "stat", "Lac-4": "exp"}

FIGURE_NUM = {}
_n = 1
for _strain, _ in STRAINS:
    FIGURE_NUM[_strain] = _n
    _n += 1
VOLCANO_FIGURE_NUM = _n


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.font.color.rgb = BLACK


def add_heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE
    run.font.color.rgb = BLACK
    run.bold = True
    return h


def add_para(doc, runs, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opts in runs:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = BODY_SIZE
        r.font.color.rgb = BLACK
        if opts.get("bold"):
            r.bold = True
        if opts.get("highlight"):
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if opts.get("italic"):
            r.italic = True
    return p


def B(text):
    return (text, {"bold": True})


def H(text):
    return (text, {"bold": True, "highlight": True})


def T(text):
    return (text, {})


def format_p(p):
    if p is None:
        return "NA"
    p = float(p)
    if p < 0.001:
        return "< 0.001"
    return f"{p:.3f}"


def load_pooled_summary(strain):
    csv_path = ANALYSIS_ROOT / strain / "fc_pooled_summary.csv"
    with open(csv_path, newline="") as f:
        rows = list(csv.DictReader(f))
    return {row["phase"]: row for row in rows}


def load_volcano_sig_counts(strain):
    csv_path = ANALYSIS_ROOT / strain / "fc_volcano_windows.csv"
    counts = {}
    with open(csv_path, newline="") as f:
        for row in csv.DictReader(f):
            key = row["phase"]
            counts.setdefault(key, {"up": 0, "down": 0, "total": 0})
            counts[key]["total"] += 1
            if row["category"].startswith("Up"):
                counts[key]["up"] += 1
            elif row["category"].startswith("Down"):
                counts[key]["down"] += 1
    return counts


def add_pooled_summary_table(doc):
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    headers = ["Strain", "Phase", "Windows", "Overall t-test p",
               "Overall Mann-Whitney U p", "Windows > +1 log2FC",
               "Windows < -1 log2FC", "Windows within ±1"]
    for cell, text in zip(header, headers):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = FONT_NAME
                r.font.size = Pt(10)
                r.font.color.rgb = BLACK
                r.bold = True

    for strain, display_name in STRAINS:
        summary = load_pooled_summary(strain)
        for phase in ["stat", "exp"]:
            row_data = summary[phase]
            row = table.add_row().cells
            values = [
                display_name, PHASE_LABEL[phase], row_data["n_windows"],
                format_p(row_data["t_test_p"]), format_p(row_data["mwu_p"]),
                row_data["n_up"], row_data["n_down"], row_data["n_ns"],
            ]
            for cell, val in zip(row, values):
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT_NAME
                        r.font.size = Pt(10)
                        r.font.color.rgb = BLACK
    doc.add_paragraph()


def add_figure(doc, image_path, fig_num, caption_text):
    doc.add_picture(str(image_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = cap.add_run(f"Figure {fig_num}. ")
    label_run.font.name = FONT_NAME
    label_run.font.size = CAPTION_SIZE
    label_run.font.color.rgb = BLACK
    label_run.bold = True
    label_run.italic = True
    text_run = cap.add_run(caption_text)
    text_run.font.name = FONT_NAME
    text_run.font.size = CAPTION_SIZE
    text_run.font.color.rgb = BLACK
    text_run.italic = True


STRAIN_INTERPRETATION = {
    "mff": [
        [T("In exponential phase (Figure 1), the overall pooled t-test is borderline "
           "non-significant ("), B("t-test p = 0.052"), T("), while the overall pooled "
           "Mann-Whitney U test is significant ("), B("p = 0.015"), T("). The two tests "
           "disagreeing at the same significance threshold is itself informative: it means "
           "any shift in read density between Meropenem-treated and untreated windows is, at "
           "most, small and inconsistently distributed, not the strong, uniform-direction "
           "shift seen in most of the other strains.")],
        [T("No single 1kb window in either phase crosses the 2-fold-change threshold "
           "(0 windows above +1 or below -1 log2FC, in both stationary and exponential "
           "phase), and the volcano plot shows no per-window significant hits. mff should be "
           "read as the flattest, least differentially affected strain in this comparison, "
           "consistent with its flat, largely featureless coverage profile already noted in "
           "the earlier replication-profile interpretation.")],
    ],
    "VU": [
        [H("VU exponential phase is the clearest fold-change signal in the whole dataset."),
         T(" The overall pooled test is highly significant in both flavors ("),
         B("t-test p < 0.001, Mann-Whitney U p < 0.001"),
         T("), and unlike every other strain/phase combination, the effect is not just a "
           "diffuse genome-wide shift: 88 windows are enriched more than 2-fold in Meropenem "
           "and 34 windows are depleted more than 2-fold (Figure 2), and these are not "
           "scattered randomly -- the enriched windows visibly cluster in a band just upstream "
           "of oriC.")],
        [T("Despite that strong pooled and spatial signal, the per-window volcano test finds "
           "zero individually significant windows in exponential phase (t-test p < 0.05 at "
           "only 2 replicates per condition). This is the clearest illustration in the dataset "
           "of why the pooled test, not the per-window test, is the one to trust for an "
           "overall significance claim: the per-window test is too underpowered to confirm "
           "any single window on its own, even where the aggregate, spatially-clustered "
           "pattern is convincing.")],
        [B("VU exponential phase is the strongest candidate for follow-up"),
         T(" among the fold-change results, particularly the enriched cluster near oriC, "
           "which is also where the earlier replication-profile analysis flagged a sharp, "
           "reproducible coverage anomaly in this strain.")],
    ],
    "AB30": [
        [T("AB30 stationary phase (Figure 3) shows a one-sided pattern: 137 windows are "
           "enriched more than 2-fold in Meropenem, and none are depleted. The overall pooled "
           "test is significant ("), B("t-test p < 0.001, Mann-Whitney U p < 0.001"),
         T("), and 87 of the 137 enriched windows also pass the stricter per-window "
           "volcano test -- the highest per-window confirmation rate of any strain/phase in "
           "this study.")],
        [T("Because this is stationary phase, where active bidirectional replication should "
           "be minimal, a one-sided, exclusively-enriched pattern is more consistent with a "
           "Meropenem-associated change in DNA accessibility or copy number at specific loci "
           "than with a replication-timing effect. Exponential phase for this strain shows no "
           "windows crossing the fold-change threshold at all, so the effect is specific to "
           "stationary phase.")],
    ],
    "AB42": [
        [T("AB42 stationary phase (Figure 4) shows the mirror-image pattern to AB30: 248 "
           "windows are depleted more than 2-fold in Meropenem, and only 2 are enriched. The "
           "overall pooled test is significant ("), B("t-test p < 0.001, Mann-Whitney U p < 0.001"),
         T("), and 6 of the 248 depleted windows also pass the per-window volcano test.")],
        [T("As with AB30, this is a stationary-phase, largely one-sided effect, arguing for a "
           "Meropenem-associated change at specific genomic regions rather than a "
           "replication-timing artifact. Exponential phase again shows no windows crossing the "
           "fold-change threshold, so -- as with AB30 -- the effect is specific to the "
           "non-replicating (stationary) condition.")],
    ],
    "Lac-4": [
        [T("Lac-4 (Figure 5) shows the most uniform result in the dataset: the overall pooled "
           "test is significant in both phases ("), B("t-test p < 0.001 in exponential, "
           "p = 0.005 in stationary"), T("), meaning there is a small, highly consistent "
           "genome-wide shift in read density between conditions, but not one single window "
           "in either phase crosses the 2-fold-change threshold.")],
        [T("This is a useful contrast with AB30/AB42: a statistically significant pooled "
           "result does not by itself imply any localized, biologically interpretable "
           "region -- it can also mean a tiny, uniform shift spread evenly across the whole "
           "genome, which is what Lac-4 shows here.")],
    ],
}


def build():
    doc = docx.Document()
    set_default_font(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run("Fold-Change Interpretation: Meropenem vs Untreated Replication Profiles")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, [
        T("Project: "), B("uea_project"), T(". Strains: "),
        B("mff, VU, AB30, AB42, Lac-4"), T(". Date: "),
        B(datetime.now().strftime("%Y-%m-%d")),
    ])
    add_para(doc, [T("Full pipeline and command log: "), B("docs/QC_Project_Log.docx"), T(".")])
    add_para(doc, [T("Underlying figures and per-window data: "),
                    B("analysis/replication_profile/<strain>/figures_fc_pooled/, "
                      "figures_fc_pooled_horizontal/, figures_fc_volcano/, and the "
                      "fc_pooled_summary.csv / fc_pooled_windows.csv / fc_volcano_windows.csv "
                      "files in the same directory"), T(".")])

    add_heading(doc, "1. What Was Built")
    add_para(doc, [T(
        "Three complementary views of the same underlying comparison were generated: for "
        "every 1kb genomic window, the mean Meropenem read-start density is divided by the "
        "mean untreated read-start density and log2-transformed (log2FC), following the "
        "conventional differential-expression style of analysis."
    )])
    add_para(doc, [B("Pooled/overall test "), T(
        "(figures_fc_pooled/): log2FC per window plotted against normalized genomic position, "
        "with vertical dashed threshold lines at log2FC = -1/0/+1. All windows for a "
        "strain/phase are pooled as the statistical sample (hundreds to thousands of windows) "
        "to run one overall t-test and one overall Mann-Whitney U test."
    )], bullet=True)
    add_para(doc, [B("Transposed / horizontal-threshold version "), T(
        "(figures_fc_pooled_horizontal/): the same data and the same statistics as above, "
        "with the axes rotated -- genomic position on x, log2FC on y, with solid horizontal "
        "threshold lines at -1/0/+1 instead of vertical dashed ones."
    )], bullet=True)
    add_para(doc, [B("Classic per-window volcano plot "), T(
        "(figures_fc_volcano/): log2FC per window (x) against -log10(per-window t-test "
        "p-value) (y), the familiar differential-expression volcano shape."
    )], bullet=True)

    add_heading(doc, "2. Why Two Different Statistical Designs")
    add_para(doc, [T(
        "Each strain/phase/treatment combination has only "), B("2 biological replicates"),
        T(" (A and B). A per-window t-test at that sample size has just 2 degrees of freedom, "
          "and a per-window Mann-Whitney U test is severely underpowered: the exact test's "
          "minimum attainable two-sided p-value is 1/3 when the four values are untied, and "
          "although R falls back to a normal approximation whenever read counts are tied "
          "(which occasionally reads a little below that nominal floor), the per-window test "
          "is fundamentally too weak to certify any single window as significant on its own.")])
    add_para(doc, [T(
        "The pooled/overall test sidesteps this by treating every window as one observation, "
        "giving each strain/phase comparison a sample size in the thousands rather than 2. "
        "This is the statistically well-powered result and the one that should be quoted as "
        "\"is there an overall difference between Meropenem and untreated\" for a given strain "
        "and phase. The per-window volcano plot is kept alongside it because it is the "
        "familiar, expected visualization for this kind of comparison and is useful for "
        "flagging which specific windows are driving a pooled result -- but a window that "
        "fails the per-window test should not be read as \"not different\"; it may simply be "
        "underpowered at n=2. Both raw and BH-adjusted per-window p-values, and the "
        "per-window Mann-Whitney U p-value, are retained in fc_volcano_windows.csv for anyone "
        "who wants to re-examine the per-window results directly."
    )])

    add_heading(doc, "3. Cross-Strain Summary: Pooled/Overall Test")
    add_pooled_summary_table(doc)
    add_para(doc, [T(
        "Read the pooled p-values as \"is there a genome-wide shift\" and the window counts "
        "as \"how much of that shift is concentrated in specific 2-fold-or-more regions, versus "
        "spread thinly across the whole genome.\" A significant pooled p-value with 0 windows "
        "above threshold (mff, Lac-4) means a small, uniform, genome-wide shift; a significant "
        "pooled p-value with hundreds of windows above threshold (AB30, AB42 stationary; VU "
        "exponential) means the effect is concentrated in specific, potentially biologically "
        "meaningful regions."
    )])

    add_heading(doc, "4. Per-Strain Interpretation")
    for strain, display_name in STRAINS:
        add_heading(doc, f"4.{STRAINS.index((strain, display_name)) + 1} Strain {display_name}", level=2)
        for line in STRAIN_INTERPRETATION[strain]:
            add_para(doc, line)
        phase = HIGHLIGHT_PHASE[strain]
        phase_prefix = "stationary" if phase == "stat" else "exponential"
        image_path = ANALYSIS_ROOT / strain / "figures_fc_pooled_horizontal" / f"{phase_prefix}_phase.png"
        if image_path.exists():
            caption = (
                f"Strain {display_name}, {PHASE_LABEL[phase].lower()} phase. Panel C: log2 "
                "fold-change (Meropenem/untreated) per 1kb window (y) against normalized "
                "genomic position (x; oriC = 0, ter = plus or minus 1). Solid horizontal lines "
                "mark 2-fold-change thresholds (log2FC = -1/0/+1). Points are colored by "
                "fold-change category. The annotation box reports the overall pooled t-test "
                "and Mann-Whitney U p-values and the window counts in each category."
            )
            add_figure(doc, image_path, FIGURE_NUM[strain], caption)
        else:
            print(f"WARNING: missing figure for {strain}: {image_path}", file=sys.stderr)

    add_heading(doc, "5. Per-Window Volcano Plot: An Illustration")
    add_para(doc, [T(
        "AB30 stationary phase is shown below (Figure 6) because it has the highest "
        "per-window confirmation rate in the dataset (87 of 137 fold-change-positive windows "
        "also pass the per-window t-test), making the classic volcano shape visible. Contrast "
        "this with VU exponential phase (Section 4.2, Figure 2), which has a stronger and more "
        "spatially clustered pooled result but zero per-window-significant hits -- a direct "
        "demonstration of the power difference between the two designs described in Section 2."
    )])
    volcano_path = ANALYSIS_ROOT / "AB30" / "figures_fc_volcano" / "stationary_phase.png"
    if volcano_path.exists():
        caption = (
            "Strain AB030, stationary phase. Panel C: classic volcano plot, log2(Meropenem/"
            "untreated) per 1kb window (x) against -log10(per-window t-test p-value) (y). "
            "Vertical lines mark log2FC = -1/+1; horizontal line marks p = 0.05. Colored points "
            "pass both the fold-change and per-window significance thresholds."
        )
        add_figure(doc, volcano_path, VOLCANO_FIGURE_NUM, caption)
    else:
        print(f"WARNING: missing figure: {volcano_path}", file=sys.stderr)

    add_heading(doc, "6. Caveats")
    add_para(doc, [T(
        "All fold-change values are computed at 1kb window resolution only. Windows with zero "
        "read coverage in any one of the four samples (Meropenem A/B, untreated A/B) are "
        "excluded from that strain/phase's comparison; the exclusion count (0 for every "
        "strain/phase in this dataset) is reported in fc_pooled_summary.csv. The 2-fold-change "
        "threshold (log2FC beyond plus or minus 1) is a conventional, not a statistically "
        "derived, cutoff. Per-window p-values, both t-test and Mann-Whitney U, should be "
        "treated as exploratory given the 2-replicate-per-condition design; the pooled/overall "
        "test is the statistically defensible claim of significance for a given strain and "
        "phase."
    )])

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
