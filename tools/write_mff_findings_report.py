#!/usr/bin/env python3
"""Generate docs/mff_Replication_Profile_Findings.docx — a formatted findings and
discussion report for the mff replication-profile analysis (separate from the
mechanical QC_Project_Log.docx step log). Times New Roman throughout; headings 14pt,
body 12pt; bold/highlight used for emphasis.
"""
from pathlib import Path
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = PROJECT_ROOT / "docs" / "mff_Replication_Profile_Findings.docx"

FONT_NAME = "Times New Roman"
HEADING_SIZE = Pt(14)
BODY_SIZE = Pt(12)


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE


def add_heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE
    run.bold = True
    return h


def add_para(doc, runs, style=None, bullet=False):
    """runs: list of (text, {bold, highlight}) tuples, or a plain string."""
    p = doc.add_paragraph(style="List Bullet" if bullet else style)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opts in runs:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = BODY_SIZE
        if opts.get("bold"):
            r.bold = True
        if opts.get("highlight"):
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if opts.get("italic"):
            r.italic = True
    return p


def B(text):
    return (text, {"bold": True})


def H(text):
    return (text, {"bold": True, "highlight": True})


def T(text):
    return (text, {})


def build():
    doc = docx.Document()
    set_default_font(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run("Replication Profile Analysis for Strain mff: Findings and Discussion")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, [
        T("Project: "), B("uea_project"), T(". Reference: "), B("CP012004.1"),
        T(" (Acinetobacter baumannii ATCC 17978-mff). Date: "),
        B(datetime.now().strftime("%Y-%m-%d")),
    ])
    add_para(doc, [T("Full step-by-step command log: "), B("docs/QC_Project_Log.docx"),
                   T(" (steps 17-23).")])

    # 1. Objective
    add_heading(doc, "1. Objective")
    add_para(doc, [T(
        "To evaluate DNA replication dynamics in strain mff by generating marker-frequency "
        "replication profiles following the methodology of Skovgaard et al. (2011, Genome "
        "Research), as a precursor step toward investigating potential chromosomal "
        "rearrangements. The expected signature of active, bidirectional chromosome "
        "replication is a tent-shaped peak in read coverage centered on the origin of "
        "replication (oriC), tapering to a minimum at the terminus (ter)."
    )])

    # 2. Methods summary
    add_heading(doc, "2. Methods Summary")
    for line in [
        [B("oriC/ter localization: "), T("independently derived using cumulative GC-skew "
           "analysis (Lobry's method) on the reference genome, and cross-checked against the "
           "user-supplied approximate coordinates.")],
        [B("Read mapping: "), T("Bowtie2, paired-end trimmed reads against the mff reference.")],
        [B("Duplicate removal: "), T("samtools sort, fixmate, and markdup. PCR duplicates were "
           "removed prior to coverage analysis, since duplicates would otherwise create false "
           "coverage spikes unrelated to real replication.")],
        [B("Marker-frequency computation: "), T("bedtools genomecov (5' read-start counting), "
           "binned into 1,000 bp and 10,000 bp windows, normalized to the genome-wide median "
           "window count (log2 ratio), with genomic position folded onto a replichore-normalized "
           "coordinate (oriC = 0, ter = plus or minus 1).")],
        [B("Visualization: "), T("R and ggplot2, six publication-format plots per the requested "
           "layout (replicate A versus B, and drug versus untreated, for stationary and "
           "exponential phase).")],
    ]:
        add_para(doc, line, bullet=True)

    # 3. Key results
    add_heading(doc, "3. Key Results")
    for line in [
        [B("Alignment rate: "), T("92.3-96.1% across all 8 samples ("), B("mean 94.1%"), T(").")],
        [B("Duplicate rate removed: "), T("12.3-16.8% ("), B("mean 15.5%"), T("), within a "
           "normal range for standard Illumina library preparation.")],
        [B("oriC/ter (GC-skew derived): "), T("oriC = "), B("1,918,501 bp"), T(", ter = "),
         B("501 bp"), T(", within approximately 500 bp of the user-supplied coordinates (oriC "
           "1,918,991, ter 521). This independently validates both the reference genome and "
           "the given coordinates.")],
    ]:
        add_para(doc, line, bullet=True)

    # 4. Central finding
    add_heading(doc, "4. Central Finding")
    add_para(doc, [H(
        "No tent-shaped replication peak was observed at oriC in any of the 8 mff samples, "
        "including exponential-phase samples where one was expected."
    )])
    add_para(doc, [T(
        "Quantitative evidence: mean log2(marker frequency) near oriC (|position| < 0.1) "
        "compared with near ter (|position| > 0.9) differed by only approximately 0.01 to "
        "0.08 log2 units, compared with the approximately 1.0 to 1.6 log2 units (a 2- to "
        "3-fold difference) reported for actively replicating cultures. This was "
        "independently cross-checked with raw "), B("samtools depth"), T(
        " at the oriC and ter regions of sample PID-2861-27 (replicate A, drug-treated, "
        "exponential phase): approximately 72 to 74x coverage uniformly across the genome, "
        "confirming that the flat marker-frequency profile is not an artifact of the "
        "windowing or binning code."
    )])
    add_para(doc, [T(
        "A sharp, localized coverage spike at normalized position approximately -0.15 "
        "appears in nearly every sample, in both replicates and both treatments. Because it "
        "is shared across independent samples rather than condition-specific, it most likely "
        "reflects a repeat or multi-copy genomic feature (for example, an IS element or "
        "rRNA operon) causing ambiguous short-read mapping, rather than a true "
        "replication-associated signal. Explicit repeat-masking, as used in the original "
        "methodology, was not implemented in this analysis and would likely remove it."
    )])

    # 5. Possible reasons
    add_heading(doc, "5. Possible Reasons for the Flat Signal")
    reasons = [
        ("PCR amplification bias during library preparation (most likely).",
         "Standard multi-cycle Illumina library preparation re-amplifies fragments before "
         "sequencing. If cycling approaches saturation, it tends to equalize the molar "
         "abundance of fragments, including the 2- to 3-fold copy-number difference that "
         "marker-frequency analysis depends on. Published marker-frequency-analysis "
         "protocols typically specify PCR-free or low-cycle library preparation for this "
         "reason."),
        ("Harvest timing not capturing synchronized exponential growth.",
         "The exponential-phase label does not by itself guarantee that cells were "
         "harvested during actively firing, well-synchronized replication. Culture density "
         "(OD) at harvest strongly affects how strong a signal is present."),
        ("Per-sample DNA or library normalization.",
         "Normalization steps that equalize input mass per sample downstream of extraction "
         "can flatten real relative abundance differences if applied at the fragment level."),
        ("Multi-fork replication.",
         "In very fast-growing cultures, overlapping rounds of replication can partially "
         "blur, though not usually fully erase, the triangular profile shape."),
        ("Unmasked repeat regions.",
         "These contribute noise and artifacts, such as the shared spike observed, but "
         "would not by themselves flatten a genuine underlying peak."),
    ]
    for header, body in reasons:
        add_para(doc, [B(header)])
        add_para(doc, [T(body)])

    # 6. Suggested strategies
    add_heading(doc, "6. Suggested Strategies Going Forward")
    for line in [
        "Verify the OD or growth-curve data behind the exponential-phase harvest label for these cultures.",
        "For any future resequencing targeted at marker-frequency analysis, consider PCR-free or low-cycle-number library preparation.",
        "Apply explicit repeat-masking (for example, self-alignment or IS-element annotation for A. baumannii) before computing marker frequency, to remove artifacts like the observed spike.",
        "Use a formal statistical test (for example, regression or changepoint analysis of the oriC-to-ter slope) rather than relying on visual inspection alone. This may reveal a small but real signal that is not obvious by eye.",
        "Increase sequencing depth and biological replicate numbers to shrink per-window noise, improving sensitivity to a weak signal.",
    ]:
        add_para(doc, [T(line)], bullet=True)

    # 7. Implications for chromosomal rearrangement analysis
    add_heading(doc, "7. Implications for Chromosomal Rearrangement Analysis")
    add_para(doc, [T(
        "Two distinct approaches to rearrangement detection are relevant here, and the flat "
        "signal affects them differently."
    )])
    add_para(doc, [B("Replication-timing-based detection "), T(
        "(comparing peak shape or asymmetry between conditions, as in the referenced "
        "methodology) "), B("is not currently supportable. "), T(
        "Without a visible baseline peak in either the drug-treated or untreated exponential "
        "samples, there is no reference shape against which to detect an alteration.")], bullet=True)
    add_para(doc, [B("Classical structural-variant or breakpoint detection "), T(
        "(discordant read pairs, split reads, local coverage discontinuities) "),
        B("is not blocked. "), T(
        "It does not depend on the genome-wide replication signal, and the existing "
        "mapped, deduplicated BAM files can be used for this directly.")], bullet=True)
    add_para(doc, [H(
        "Conclusion: chromosomal rearrangement investigation can proceed, but through "
        "structural-variant or breakpoint calling rather than replication-peak-shape comparison."
    )])

    # 8. Long-read recommendation
    add_heading(doc, "8. Recommendation on Long-Read Sequencing")
    add_para(doc, [T(
        "Short reads cannot uniquely span repeats or IS elements, which are the plausible "
        "cause of the shared coverage spike observed in every sample, and these are exactly "
        "the sequence contexts where bacterial chromosomal rearrangements typically occur "
        "through homologous recombination. A single long read spanning a rearrangement "
        "junction provides direct, unambiguous evidence, independent of coverage-based "
        "statistics."
    )])
    add_para(doc, [H(
        "Recommendation: use a long-read assembly and whole-genome alignment against the "
        "reference as the primary method for chromosomal rearrangement detection in this "
        "project (long_read_data is already available for these strains), with the "
        "short-read marker-frequency analysis retained as secondary, supplementary evidence "
        "rather than the primary method."
    )])

    # 9. Conclusion
    add_heading(doc, "9. Conclusion and Next Steps")
    add_para(doc, [T(
        "The mapping and marker-frequency pipeline is validated and working correctly. The "
        "absence of a replication peak in strain mff appears to be a genuine finding, most "
        "plausibly attributable to library-preparation PCR bias or harvest-timing effects "
        "rather than a bioinformatics error. This does not prevent chromosomal rearrangement "
        "investigation from proceeding. It redirects that investigation toward "
        "structural-variant calling from the existing short-read alignments and, preferably, "
        "toward a long-read-based assembly comparison as the definitive method."
    )])

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
