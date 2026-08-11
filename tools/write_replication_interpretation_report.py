#!/usr/bin/env python3
"""Generate docs/Replication_Profile_Interpretation.docx: a formatted, per-strain
interpretation of the replication-profile figures and Ori:Ter ratios for all five
strains. Times New Roman throughout; headings 14pt, body 12pt; bold/highlight for
emphasis.
"""
import csv
import sys
from pathlib import Path
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = PROJECT_ROOT / "docs" / "Replication_Profile_Interpretation.docx"
ANALYSIS_ROOT = PROJECT_ROOT / "analysis" / "replication_profile"

FONT_NAME = "Times New Roman"
HEADING_SIZE = Pt(14)
BODY_SIZE = Pt(12)
CAPTION_SIZE = Pt(11)
BLACK = RGBColor(0, 0, 0)

STRAINS = [
    ("mff", "ATCC 17978-mff"),
    ("VU", "ATCC 17978-VU"),
    ("AB30", "AB030"),
    ("AB42", "AB042"),
    ("Lac-4", "LAC-4"),
]

# Figure numbers are fixed by strain order and phase (stationary, then exponential),
# so both the in-text citations and the Figures section stay in sync.
FIGURE_NUM = {}
_n = 1
for _strain, _ in STRAINS:
    FIGURE_NUM[(_strain, "stat")] = _n
    _n += 1
    FIGURE_NUM[(_strain, "exp")] = _n
    _n += 1


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.font.color.rgb = BLACK


def add_heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = HEADING_SIZE
    run.font.color.rgb = BLACK
    run.bold = True
    return h


def add_para(doc, runs, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opts in runs:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = BODY_SIZE
        r.font.color.rgb = BLACK
        if opts.get("bold"):
            r.bold = True
        if opts.get("highlight"):
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if opts.get("italic"):
            r.italic = True
    return p


def B(text):
    return (text, {"bold": True})


def H(text):
    return (text, {"bold": True, "highlight": True})


def T(text):
    return (text, {})


def add_ratio_table(doc, strain):
    csv_path = ANALYSIS_ROOT / strain / "ori_ter_ratios.csv"
    with open(csv_path, newline="") as f:
        rows = list(csv.DictReader(f))

    series_order = ["drug_A", "drug_B", "ND_A", "ND_B", "drug_avg", "ND_avg"]
    series_label = {
        "drug_A": "Drug-treated, replicate A",
        "drug_B": "Drug-treated, replicate B",
        "ND_A": "Untreated, replicate A",
        "ND_B": "Untreated, replicate B",
        "drug_avg": "Drug-treated (replicate-averaged)",
        "ND_avg": "Untreated (replicate-averaged)",
    }
    by_series_phase = {}
    for row in rows:
        by_series_phase[(row["series"], row["phase"])] = float(row["ori_ter_ratio"])

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Series"
    header[1].text = "Stationary phase"
    header[2].text = "Exponential phase"
    for cell in header:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = FONT_NAME
                r.font.size = BODY_SIZE
                r.font.color.rgb = BLACK
                r.bold = True

    for series in series_order:
        stat_val = by_series_phase.get((series, "stat"))
        exp_val = by_series_phase.get((series, "exp"))
        row = table.add_row().cells
        row[0].text = series_label[series]
        row[1].text = f"{stat_val:.2f}" if stat_val is not None else "-"
        row[2].text = f"{exp_val:.2f}" if exp_val is not None else "-"
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = BODY_SIZE
                    r.font.color.rgb = BLACK
    doc.add_paragraph()


STRAIN_INTERPRETATION = {
    "mff": [
        [T("Ori:Ter ratios cluster tightly around 1.0 in both phases (stationary 0.94-0.97, "
           "Figure 1; exponential 1.00-1.06, Figure 2). There is no meaningful oriC-proximal "
           "excess in either growth condition.")],
        [T("The small shift from just below 1.0 in stationary phase (Figure 1) to just above "
           "1.0 in exponential phase (Figure 2) is in the biologically expected direction "
           "(more replication activity when actively growing), but the magnitude is "
           "negligible. This matches the earlier finding for this strain: a flat, largely "
           "featureless coverage profile, most plausibly due to library-preparation PCR "
           "amplification masking the true replication signal rather than a mapping or "
           "analysis error.")],
        [H("mff is not informative for replication-timing-based rearrangement detection on "
           "its own."), T(" Structural-variant calling from the existing alignments, or a "
           "long-read comparison, remains the more promising route for this strain.")],
    ],
    "VU": [
        [T("Drug-treated exponential cultures show a moderate, reproducible Ori:Ter elevation "
           "(replicate A 1.05, replicate B 1.30, averaged 1.16; Figure 4), while untreated "
           "exponential cultures stay close to parity (1.01-1.04, Figure 4). Stationary-phase "
           "ratios are close to or slightly below 1.0 in both treatments (Figure 3).")],
        [H("The standout feature in this strain is not the ratio at all: every sample, in both "
           "phases and both treatments (Figures 3 and 4), shows a sharp, localized coverage "
           "collapse (down to roughly -3 to -4.6 on the log2 scale) at a fixed position just "
           "past oriC."),
         T(" Because it appears identically across independent biological replicates and "
           "treatments, it is very unlikely to be noise. It most plausibly reflects a real "
           "structural feature, such as a deletion, an insertion element, or a region where "
           "the sequenced isolates diverge from the deposited reference assembly, rather than "
           "a replication-timing effect.")],
        [B("This localized anomaly is the most actionable finding across all five strains for "
           "the chromosomal-rearrangement objective."), T(" It is not fully captured by the "
           "Ori:Ter ratio, since the ratio summarizes broad regions near oriC and ter rather "
           "than a single narrow position, and is worth direct follow-up: manual inspection of "
           "reads at that locus, or confirmation against the long-read data for this strain.")],
    ],
    "AB30": [
        [T("Untreated cultures show a consistent oriC-proximal excess in both phases "
           "(stationary 1.11-1.22, Figure 5; exponential 1.11-1.13, Figure 6), more pronounced "
           "in stationary phase than would be expected if it were purely a replication-timing "
           "effect. Drug-treated cultures start close to parity in stationary phase (0.96-0.99, "
           "Figure 5) and rise toward untreated levels in exponential phase (1.03-1.13, "
           "Figure 6).")],
        [T("The pattern is modest in magnitude but reproducible across both biological "
           "replicates in every condition, which argues against it being a one-off artifact. "
           "The fact that untreated stationary cultures already show elevation is notable and "
           "suggests this may partly reflect a stable strain-level feature near oriC rather "
           "than purely active replication.")],
    ],
    "AB42": [
        [H("AB042 shows the most consistent oriC-proximal excess of any strain in this study, "
           "present in both growth phases and both treatments (Figures 7 and 8)."),
         T(" Ratios range from 1.03 to 1.19 in exponential phase (Figure 8) and 1.08 to 1.45 "
           "in stationary phase (Figure 7), with the single highest value in the dataset (1.45) "
           "in stationary, drug-treated replicate A.")],
        [T("Because this elevation persists in stationary phase, where active chromosome "
           "replication should be minimal, it is more consistent with a stable, "
           "replication-independent copy-number feature near oriC (for example, a partial "
           "gene duplication already present in the genome) than with an actively firing "
           "replication fork.")],
        [B("Of the five strains, AB042 is the strongest candidate for direct follow-up"),
         T(" given how consistently the elevation appears across replicates, treatments, and "
           "growth phases.")],
    ],
    "Lac-4": [
        [T("Lac-4 shows no oriC-proximal excess in either phase (Figures 9 and 10); most "
           "series sit at or slightly below 1.0 (range 0.89-1.05), if anything trending toward "
           "a mild excess near ter, which is the opposite of the expected replication-associated "
           "direction.")],
        [T("Under drug treatment in exponential phase (Figure 10), the two biological "
           "replicates disagree in direction (replicate A 0.89, replicate B 1.05), which "
           "further weakens confidence in any signal for this strain and condition.")],
        [T("Lac-4 should be treated as a negative or uninformative result for this analysis, "
           "not as a technical failure; the mapping and coverage computation succeeded "
           "cleanly for all 8 samples (see the QC log), the data simply do not show a "
           "replication-associated pattern.")],
    ],
}


def add_figure(doc, image_path, fig_num, caption_text):
    doc.add_picture(str(image_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = cap.add_run(f"Figure {fig_num}. ")
    label_run.font.name = FONT_NAME
    label_run.font.size = CAPTION_SIZE
    label_run.font.color.rgb = BLACK
    label_run.bold = True
    label_run.italic = True
    text_run = cap.add_run(caption_text)
    text_run.font.name = FONT_NAME
    text_run.font.size = CAPTION_SIZE
    text_run.font.color.rgb = BLACK
    text_run.italic = True


def add_figures_section(doc):
    add_heading(doc, "6. Figures")
    add_para(doc, [T(
        "Figures are ordered by strain, stationary phase followed by exponential phase, and "
        "are cited by number in the interpretation above."
    )])

    first = True
    for strain, display_name in STRAINS:
        for phase_key, phase_label, file_prefix in [
            ("stat", "Stationary", "stationary"),
            ("exp", "Exponential", "exponential"),
        ]:
            fig_num = FIGURE_NUM[(strain, phase_key)]
            image_path = ANALYSIS_ROOT / strain / "figures" / f"{file_prefix}_phase.png"
            if not image_path.exists():
                print(f"WARNING: missing figure for {strain} {phase_key}: {image_path}", file=sys.stderr)
                continue
            if not first:
                doc.add_page_break()
            first = False
            caption = (
                f"Strain {display_name}, {phase_label.lower()} phase marker-frequency "
                "replication profile. Panel A: drug-treated, biological replicate A versus B. "
                "Panel B: untreated (no drug), replicate A versus B. Panel C: drug-treated "
                "versus untreated, replicate-averaged. oriC = 0, ter = plus or minus 1 on the "
                "x-axis. Ori:Ter is the median marker-frequency ratio between windows near "
                "oriC and near ter, annotated on each panel."
            )
            add_figure(doc, image_path, fig_num, caption)


def build():
    doc = docx.Document()
    set_default_font(doc)

    title = doc.add_heading("", level=0)
    run = title.add_run("Replication Profile Interpretation: Five Acinetobacter baumannii Strains")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, [
        T("Project: "), B("uea_project"), T(". Strains: "),
        B("mff, VU, AB30, AB42, Lac-4"), T(". Date: "),
        B(datetime.now().strftime("%Y-%m-%d")),
    ])
    add_para(doc, [T("Full pipeline and command log: "), B("docs/QC_Project_Log.docx"), T(".")])
    add_para(doc, [T("Figures referenced here: "), B("analysis/replication_profile/<strain>/figures/"), T(".")])

    add_heading(doc, "1. What Is Plotted")
    add_para(doc, [T(
        "Each figure shows log2(marker frequency) (read-start count relative to the "
        "genome-wide median, at 1,000 bp and 10,000 bp window resolution) against genomic "
        "position, normalized so the origin of replication (oriC) sits at 0 and the terminus "
        "(ter) sits at plus or minus 1. All panels, across every strain and phase, use the "
        "same fixed y-axis (-5 to 3) so figures are directly comparable to one another."
    )])

    add_heading(doc, "2. The Ori:Ter Ratio")
    add_para(doc, [T(
        "Each panel is annotated with an Ori:Ter ratio: the median marker frequency in "
        "windows within 10% of oriC, divided by the median marker frequency in windows "
        "within 10% of ter, converted back from the log2 scale. A ratio of 1.0 means no "
        "difference between the two regions; a ratio above 1.0 means oriC-proximal windows "
        "have more read coverage than ter-proximal windows, the signature expected during "
        "active bidirectional replication. This mirrors the ori/ter ratio reported in "
        "Skovgaard et al. (2011), and replaces the regression-based statistics used in the "
        "previous version of these figures with a single, more directly interpretable number."
    )])

    add_heading(doc, "3. Per-Strain Interpretation")
    for strain, display_name in STRAINS:
        add_heading(doc, f"3.{STRAINS.index((strain, display_name)) + 1} Strain {display_name}", level=2)
        add_ratio_table(doc, strain)
        for line in STRAIN_INTERPRETATION[strain]:
            add_para(doc, line)

    add_heading(doc, "4. Cross-Strain Summary")
    add_para(doc, [B("AB042 "), T("shows the strongest and most consistent oriC-proximal excess, present in both growth phases.")], bullet=True)
    add_para(doc, [B("AB30 "), T("shows a modest, reproducible oriC-proximal excess, mainly in untreated cultures.")], bullet=True)
    add_para(doc, [B("VU "), T("shows little ratio elevation except under drug treatment in exponential phase, but carries a separate, striking, localized coverage anomaly near oriC that is unrelated to the ratio metric and is the most actionable single finding in this dataset.")], bullet=True)
    add_para(doc, [B("mff and Lac-4 "), T("show no meaningful oriC-proximal excess in either phase and should be treated as negative results for this analysis.")], bullet=True)

    add_heading(doc, "5. Caveats")
    add_para(doc, [T(
        "The Ori:Ter ratio is computed from windows within a fixed 10% zone of oriC and ter; "
        "a localized feature that happens to sit near the edge of that zone, as in VU, can "
        "influence the ratio without being fully captured by it. Explicit repeat-masking, as "
        "used in the reference methodology, was not applied here, so isolated high- or "
        "low-coverage windows from repetitive or multi-copy sequence can still appear in the "
        "raw scatter even where the ratio itself looks unremarkable. Ratios near 1.0 in "
        "stationary phase are the expected baseline, not a sign of a failed analysis."
    )])

    add_figures_section(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
